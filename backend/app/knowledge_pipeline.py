from __future__ import annotations

import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from pypdf import PdfReader

from .bailian import BailianSyncClient
from .config import AppConfig
from .types import EvalCase, FaqEntry, IngestReport, KnowledgeChunk, SourceDocument
from .utils import (
    chunked_paragraphs,
    clean_text,
    excerpt_text,
    md5_file,
    sha256_text,
    slugify,
    unique_by_key,
    write_json,
    write_jsonl,
)


MATERIAL_KEYWORDS = {
    "植鞣革": ["植鞣革", "手柄变黑", "本色皮"],
    "小羊皮": ["小羊皮", "羊皮"],
    "头层牛皮": ["头层牛皮", "牛皮"],
    "翻毛皮": ["翻毛皮", "反绒皮", "磨砂皮"],
    "漆皮": ["漆皮"],
    "涂层帆布": ["老花帆布", "涂层帆布", "老花"],
    "五金": ["拉链", "五金", "金属"],
}

DAMAGE_KEYWORDS = {
    "发黑污渍": ["发黑", "污渍", "变黑"],
    "干裂发硬": ["干裂", "发硬", "淋雨"],
    "划痕": ["划痕", "刮痕"],
    "变形塌陷": ["变形", "软塌", "塌陷"],
    "染色": ["染色", "串色"],
    "油渍": ["油渍"],
    "磨损破洞": ["磨损", "破洞", "掉色"],
    "霉变": ["发霉", "霉斑", "白毛"],
    "边油开裂": ["边油", "裂开", "脱落"],
    "锈蚀卡顿": ["生锈", "卡顿", "不顺滑"],
}

RISK_KEYWORDS = {
    "high": ["暴晒", "吹风机", "酒精", "卸甲水", "水洗", "强溶剂"],
    "medium": ["补色", "修补", "加热", "边油"],
}


class KnowledgePipeline:
    def __init__(self, config: AppConfig) -> None:
        self.config = config

    def ingest(self, sync_cloud: bool = True) -> dict[str, Any]:
        self._ensure_directories()
        raw_sources = self._load_raw_sources()
        normalized_sources = [self._normalize_source(source) for source in raw_sources]
        chunks = self._build_chunks(normalized_sources)
        faq_entries = self._build_faq_entries(chunks)
        eval_cases = self._build_eval_cases(faq_entries)

        docs_bundle = self._write_docs_bundle(normalized_sources)
        faq_bundle = self._write_faq_bundle(faq_entries)
        manual_import = self._write_bailian_import_checklist(raw_sources)
        sync_result = BailianSyncClient(self.config).sync_bundles(docs_bundle, faq_bundle) if sync_cloud else None

        report = IngestReport(
            generated_at=datetime.now(timezone.utc).isoformat(),
            source_count=len(raw_sources),
            normalized_count=len(normalized_sources),
            chunk_count=len(chunks),
            faq_count=len(faq_entries),
            eval_count=len(eval_cases),
            sources=[source.to_dict() for source in normalized_sources],
            sync=(sync_result.to_dict() if sync_result else {"status": "skipped", "detail": "Cloud sync disabled for this run."}),
            manual_import=manual_import,
        )

        write_jsonl(self.config.chunk_manifest_path, [chunk.to_dict() for chunk in chunks])
        write_jsonl(self.config.faq_manifest_path, [faq.to_dict() for faq in faq_entries])
        write_jsonl(self.config.eval_manifest_path, [case.to_dict() for case in eval_cases])
        write_json(self.config.ingest_report_path, report.to_dict())
        return report.to_dict()

    def _ensure_directories(self) -> None:
        for path in [
            self.config.raw_dir,
            self.config.normalized_dir,
            self.config.chunks_dir,
            self.config.faq_dir,
            self.config.eval_dir,
            self.config.manifest_dir,
        ]:
            path.mkdir(parents=True, exist_ok=True)

    def _load_raw_sources(self) -> list[SourceDocument]:
        sources: list[SourceDocument] = []
        for path in sorted(self.config.raw_dir.glob("*")):
            if not path.is_file() or path.name.startswith("."):
                continue
            content = self._read_file(path)
            if not content.strip():
                continue
            title = self._infer_title(path, content)
            text = clean_text(content)
            metadata = self._extract_metadata(text=text, title=title, source_path=path)
            sources.append(
                SourceDocument(
                    source_id=slugify(path.stem),
                    source_path=str(path.relative_to(self.config.project_root)),
                    title=title,
                    content=text,
                    kind=path.suffix.lstrip("."),
                    hash_value=md5_file(path),
                    metadata=metadata,
                )
            )
        return sources

    def _read_file(self, path: Path) -> str:
        suffix = path.suffix.lower()
        if suffix in {".md", ".txt"}:
            return path.read_text(encoding="utf-8")
        if suffix == ".docx":
            document = DocxDocument(str(path))
            return "\n".join(paragraph.text for paragraph in document.paragraphs)
        if suffix == ".pdf":
            reader = PdfReader(str(path))
            return "\n".join((page.extract_text() or "") for page in reader.pages)
        return ""

    def _infer_title(self, path: Path, content: str) -> str:
        first_heading = re.search(r"^#\s+(.+)$", content, re.M)
        if first_heading:
            return clean_text(first_heading.group(1))
        return path.stem.replace("-", " ").replace("_", " ")

    def _normalize_source(self, source: SourceDocument) -> SourceDocument:
        frontmatter_lines = [
            f"# {source.title}",
            "",
            f"- source_path: {source.source_path}",
            f"- material: {', '.join(source.metadata.get('materials', [])) or '未标注'}",
            f"- damage_type: {', '.join(source.metadata.get('damage_types', [])) or '未标注'}",
            f"- risk_level: {source.metadata.get('risk_level', 'low')}",
            "",
            source.content,
            "",
        ]
        normalized_content = "\n".join(frontmatter_lines)
        normalized_source = SourceDocument(
            source_id=source.source_id,
            source_path=source.source_path,
            title=source.title,
            content=normalized_content,
            kind=source.kind,
            hash_value=source.hash_value,
            metadata=source.metadata,
        )
        normalized_path = self.config.normalized_dir / f"{source.source_id}.md"
        normalized_path.write_text(normalized_content, encoding="utf-8")
        return normalized_source

    def _build_chunks(self, normalized_sources: list[SourceDocument]) -> list[KnowledgeChunk]:
        chunks: list[KnowledgeChunk] = []
        for source in normalized_sources:
            for index, paragraph_block in enumerate(chunked_paragraphs(source.content)):
                chunk_id = f"{source.source_id}-chunk-{index + 1}"
                chunk_metadata = {
                    **source.metadata,
                    "source_path": source.source_path,
                    "chunk_index": index + 1,
                    "store": "docs",
                }
                chunk = KnowledgeChunk(
                    chunk_id=chunk_id,
                    source_id=source.source_id,
                    title=source.title,
                    content=paragraph_block,
                    excerpt=excerpt_text(paragraph_block),
                    metadata=chunk_metadata,
                )
                chunk_path = self.config.chunks_dir / f"{chunk_id}.md"
                chunk_path.write_text(paragraph_block, encoding="utf-8")
                chunks.append(chunk)
        return chunks

    def _build_faq_entries(self, chunks: list[KnowledgeChunk]) -> list[FaqEntry]:
        faq_entries = self._parse_seed_faq_entries()
        generated: list[FaqEntry] = []
        for chunk in chunks:
            materials = chunk.metadata.get("materials") or ["皮具"]
            damages = chunk.metadata.get("damage_types") or ["日常养护"]
            primary_material = materials[0]
            primary_damage = damages[0]
            question = f"{primary_material}出现{primary_damage}时应该怎么处理？"
            answer = self._chunk_to_answer(chunk)
            generated.append(
                FaqEntry(
                    faq_id=f"faq-{sha256_text(question + chunk.chunk_id)[:12]}",
                    question=question,
                    answer=answer,
                    source_id=chunk.source_id,
                    title=chunk.title,
                    metadata={**chunk.metadata, "store": "faq"},
                )
            )
        deduped = unique_by_key([faq.to_dict() for faq in faq_entries + generated], "question")
        return [FaqEntry(**item) for item in deduped]

    def _parse_seed_faq_entries(self) -> list[FaqEntry]:
        if not self.config.qa_seed_path.exists():
            return []
        text = self.config.qa_seed_path.read_text(encoding="utf-8")
        pattern = re.compile(r"\*\*Q:\s*(.+?)\*\*\nA:\s*(.+?)(?=\n\*\*Q:|\Z)", re.S)
        entries: list[FaqEntry] = []
        for index, match in enumerate(pattern.finditer(text), start=1):
            question = clean_text(match.group(1))
            answer = clean_text(match.group(2))
            metadata = self._extract_metadata(text=f"{question}\n{answer}", title="FAQ Seed", source_path=self.config.qa_seed_path)
            entries.append(
                FaqEntry(
                    faq_id=f"seed-faq-{index}",
                    question=question,
                    answer=answer,
                    source_id="seed-faq",
                    title="皮具修复核心知识点",
                    metadata={**metadata, "source_path": str(self.config.qa_seed_path.relative_to(self.config.project_root)), "store": "faq"},
                )
            )
        return entries

    def _build_eval_cases(self, faq_entries: list[FaqEntry]) -> list[EvalCase]:
        cases: list[EvalCase] = []
        for index, faq in enumerate(faq_entries[:12], start=1):
            keywords = []
            keywords.extend(faq.metadata.get("materials", [])[:1])
            keywords.extend(faq.metadata.get("damage_types", [])[:1])
            keywords.extend(["注意事项", "操作步骤"])
            cases.append(
                EvalCase(
                    case_id=f"eval-{index}",
                    question=faq.question,
                    expected_keywords=[keyword for keyword in keywords if keyword],
                    title=faq.title,
                    metadata=faq.metadata,
                )
            )
        return cases

    def _write_docs_bundle(self, normalized_sources: list[SourceDocument]) -> Path:
        lines = ["# Leather Care Documents", ""]
        for source in normalized_sources:
            lines.extend(
                [
                    f"## {source.title}",
                    f"- source_path: {source.source_path}",
                    f"- material: {', '.join(source.metadata.get('materials', [])) or '未标注'}",
                    f"- damage_type: {', '.join(source.metadata.get('damage_types', [])) or '未标注'}",
                    "",
                    source.content,
                    "",
                ]
            )
        self.config.generated_docs_bundle.write_text("\n".join(lines), encoding="utf-8")
        return self.config.generated_docs_bundle

    def _write_faq_bundle(self, faq_entries: list[FaqEntry]) -> Path:
        lines = ["# Leather Care FAQ Bundle", ""]
        for faq in faq_entries:
            lines.extend(
                [
                    f"## {faq.question}",
                    f"- source_id: {faq.source_id}",
                    f"- title: {faq.title}",
                    "",
                    faq.answer,
                    "",
                ]
            )
        self.config.generated_faq_bundle.write_text("\n".join(lines), encoding="utf-8")
        return self.config.generated_faq_bundle

    def _write_bailian_import_checklist(self, raw_sources: list[SourceDocument]) -> dict[str, Any]:
        source_paths = [source.source_path for source in raw_sources]
        checklist_lines = [
            "# 百炼知识库导入清单",
            "",
            f"- 目标知识库 ID: {self.config.docs_kb_id or '未配置'}",
            "- 运行时检索模式: 本地 LangChain 检索",
            "- 百炼角色: 云端知识库归档与答辩展示",
            "- 上传策略: 手动上传独立 Markdown 文件",
            f"- 建议上传数量: {len(source_paths)}",
            f"- 不建议上传的合并包: {self.config.generated_docs_bundle.relative_to(self.config.project_root)}",
            "",
            "## 上传步骤",
            "",
            "1. 打开百炼控制台中的目标知识库。",
            "2. 逐个上传下列独立 Markdown 文件。",
            "3. 不要使用 docs_kb_bundle.md 作为主知识库文件。",
            "4. 上传完成后，确认云端文档数与下列清单一致。",
            "",
            "## 推荐上传文件",
            "",
        ]
        checklist_lines.extend(
            f"- {source.source_path} | {source.title}"
            for source in raw_sources
        )
        checklist_lines.extend(
            [
                "",
                "## 说明",
                "",
                "- FAQ 合并包保留为本地检索和论文展示材料。",
                "- 若目标知识库 ID 未配置，可先在 .env 中填写 BAILIAN_DOCS_KB_ID 后重新 ingest。",
            ]
        )
        self.config.bailian_import_checklist_path.write_text("\n".join(checklist_lines), encoding="utf-8")
        return {
            "strategy": "manual_upload",
            "target_docs_kb_id": self.config.docs_kb_id,
            "recommended_file_count": len(source_paths),
            "recommended_files": source_paths,
            "checklist_path": str(self.config.bailian_import_checklist_path.relative_to(self.config.project_root)),
            "avoid_bundle_path": str(self.config.generated_docs_bundle.relative_to(self.config.project_root)),
            "runtime_retrieval_mode": "local_langchain",
        }

    def _chunk_to_answer(self, chunk: KnowledgeChunk) -> str:
        materials = ", ".join(chunk.metadata.get("materials", []) or ["皮具"])
        damages = ", ".join(chunk.metadata.get("damage_types", []) or ["常见问题"])
        return "\n".join(
            [
                "### 适用判断",
                f"以下建议适用于{materials}相关的{damages}场景，若出现大面积开裂、渗色或结构损坏，应优先咨询专业修复。",
                "",
                "### 所需工具",
                "请根据材料准备软布、中性清洁用品、护理油或专用喷剂，并避免强溶剂。",
                "",
                "### 操作步骤",
                chunk.excerpt,
                "",
                "### 注意事项",
                f"风险等级：{chunk.metadata.get('risk_level', 'low')}。先在不显眼处测试，再小范围处理。",
                "",
                "### 何时送修",
                "出现严重掉色、破洞、结构变形或多次处理仍无改善时，应送专业门店。",
                "",
                "### 参考来源",
                f"{chunk.title} / {chunk.metadata.get('source_path', '')}",
            ]
        )

    def _extract_metadata(self, text: str, title: str, source_path: Path) -> dict[str, Any]:
        materials = [name for name, keywords in MATERIAL_KEYWORDS.items() if any(keyword in text for keyword in keywords)]
        damage_types = [name for name, keywords in DAMAGE_KEYWORDS.items() if any(keyword in text for keyword in keywords)]

        risk_level = "low"
        for level, keywords in RISK_KEYWORDS.items():
            if any(keyword in text for keyword in keywords):
                risk_level = level
                break

        return {
            "materials": materials,
            "damage_types": damage_types,
            "risk_level": risk_level,
            "source_name": title,
            "source_file": source_path.name,
        }
